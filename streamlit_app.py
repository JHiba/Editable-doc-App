import io
import os
import json
import base64
import tempfile
import streamlit as st
from PIL import Image
import mammoth  # Added for live preview
from fpdf import FPDF
import fitz  # PyMuPDF

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from groq import Groq

# -------------------------
# Page Config & Constants
# -------------------------
st.set_page_config(page_title="Image to Formatted DOCX ", page_icon="📄", layout="centered")

# -------------------------
# Groq + Vision helpers
# -------------------------
def get_groq_client() -> Groq:
    api_key = os.getenv("GROQ_API_KEY") or st.secrets.get("GROQ_API_KEY", None)
    if not api_key:
        raise RuntimeError("Missing GROQ_API_KEY. Set it in environment variables or Streamlit secrets.")
    return Groq(api_key=api_key)

def process_uploaded_files(uploaded_files):
    images = []
    for f in uploaded_files:
        if f.name.lower().endswith(".pdf"):
            try:
                # Convert PDF pages to PIL images using PyMuPDF (no external dependencies needed)
                pdf_bytes = f.read()
                pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc.load_page(page_num)
                    pix = page.get_pixmap(dpi=200) # High quality render
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    images.append(img)
            except Exception as e:
                st.error(f"Failed to read PDF {f.name}: {str(e)}")
                st.stop()
        else:
            images.append(Image.open(f))
    return images

def pil_to_data_url(pil_img: Image.Image) -> str:
    if pil_img.mode != "RGB":
        pil_img = pil_img.convert("RGB")
    buf = io.BytesIO()
    pil_img.save(buf, format="JPEG", quality=90)
    b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    return f"data:image/jpeg;base64,{b64}"

SYSTEM_PROMPT = """
You are a document conversion engine.
Your job: extract text, formatting, and diagram regions from a single-page English document image.

Return ONLY valid JSON.
Formatting to detect:
- paragraphs (separate blocks of text)
- alignment for each paragraph: left, center, right, justify
- bold and italic within text runs
- diagrams: if there are any hand-drawn or printed diagrams, graphs, illustrations, OR chemical/mathematical equations, YOU MUST define their bounding box in the "diagrams" array. DO NOT extract complex chemical equations as text paragraphs. Instead, map their bounding box so they can be cropped!

Rules:
- Bounding box 'bbox' must be [x_min, y_min, x_max, y_max] representing percentages (0 to 100) of the image width and height.
- Example: [10.5, 20.0, 90.0, 45.5] means x starts at 10.5%, y starts at 20.0%, etc.
- Preserve the reading order top-to-bottom.
"""

USER_PROMPT = """
Extract the document into this JSON shape:
{
  "paragraphs": [
    {
      "alignment": "left|center|right|justify",
      "runs": [
        {"text": "string", "bold": true|false, "italic": true|false}
      ]
    }
  ],
  "diagrams": [
    {
      "bbox": [x_min, y_min, x_max, y_max]
    }
  ]
}
"""

def groq_extract_layout(pil_img: Image.Image, model_id: str) -> dict:
    client = get_groq_client()
    data_url = pil_to_data_url(pil_img)

    for attempt in range(2):
        temperature = 0.2 if attempt == 0 else 0.0
        completion = client.chat.completions.create(
            model=model_id,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT.strip()},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": USER_PROMPT.strip()},
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                },
            ],
            response_format={"type": "json_object"},
            temperature=temperature,
            max_completion_tokens=2048,
        )

        content = completion.choices[0].message.content
        try:
            obj = json.loads(content)
            if "paragraphs" not in obj or not isinstance(obj["paragraphs"], list):
                raise ValueError("Bad JSON shape")
            return obj
        except Exception:
            if attempt == 1:
                raise RuntimeError("Groq returned invalid JSON.")
    return {"paragraphs": []}

AGENT_SYSTEM_PROMPT = """
You are an intelligent document analysis agent.
Your task is to perceive the uploaded document image and make initial assessments.
Return ONLY valid JSON.
Identify:
1. 'has_diagrams': true if the document contains non-text hand-drawn/printed diagrams, graphs, illustrations, OR complex mathematical/chemical equations.
2. 'has_pii': true if the document contains Personally Identifiable Information (PII) like names, phone numbers, addresses, SSNs, or IDs.
3. 'reasoning': A brief 1-2 sentence explanation of your findings.
"""

AGENT_USER_PROMPT = """
Analyze this document and return the assessment in this JSON shape:
{
  "has_diagrams": true|false,
  "has_pii": true|false,
  "reasoning": "string"
}
"""

def groq_agent_perceive(pil_img: Image.Image, model_id: str) -> dict:
    client = get_groq_client()
    data_url = pil_to_data_url(pil_img)
    
    completion = client.chat.completions.create(
        model=model_id,
        messages=[
            {"role": "system", "content": AGENT_SYSTEM_PROMPT.strip()},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": AGENT_USER_PROMPT.strip()},
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            },
        ],
        response_format={"type": "json_object"},
        temperature=0.1,
        max_completion_tokens=500,
    )
    
    content = completion.choices[0].message.content
    try:
        return json.loads(content)
    except Exception:
        return {"has_diagrams": False, "has_pii": False, "reasoning": "Failed to parse agent perception."}


# -------------------------
# DOCX building
# -------------------------
def add_image_to_doc(doc: Document, pil_img: Image.Image, width_in=6.2):
    buf = io.BytesIO()
    pil_img.save(buf, format="PNG")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width_in))

def alignment_to_docx(alignment: str):
    a = (alignment or "").lower().strip()
    if a == "center": return WD_ALIGN_PARAGRAPH.CENTER
    if a == "right": return WD_ALIGN_PARAGRAPH.RIGHT
    if a == "justify": return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT

def process_document(images, model_id: str, include_page_image: bool):
    doc = Document()
    doc.add_heading("Image to Formatted DOCX", level=1)
    
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    cropped_count = 0
    
    for idx, pil_img in enumerate(images, start=1):
        layout = groq_extract_layout(pil_img, model_id=model_id)
        
        doc.add_heading(f"Page {idx}", level=2)
        pdf.add_page()
        pdf.set_font("Helvetica", style="B", size=16)
        pdf.cell(0, 10, text=f"Page {idx}", new_x="LMARGIN", new_y="NEXT")
        
        if include_page_image:
            add_image_to_doc(doc, pil_img)
            doc.add_paragraph("")
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                pil_img.save(tmp.name)
                pdf.image(tmp.name, w=180)
            os.unlink(tmp.name)
            
        elif layout.get("diagrams"):
            # Agent intelligently crops diagrams
            for d in layout.get("diagrams", []):
                bbox = d.get("bbox")
                if bbox and len(bbox) == 4:
                    try:
                        img_width, img_height = pil_img.size
                        x1 = int((bbox[0] / 100.0) * img_width)
                        y1 = int((bbox[1] / 100.0) * img_height)
                        x2 = int((bbox[2] / 100.0) * img_width)
                        y2 = int((bbox[3] / 100.0) * img_height)
                        
                        # Ensure bounds
                        x1, y1 = max(0, x1), max(0, y1)
                        x2, y2 = min(img_width, x2), min(img_height, y2)
                        
                        if x2 > x1 and y2 > y1:
                            cropped = pil_img.crop((x1, y1, x2, y2))
                            add_image_to_doc(doc, cropped, width_in=5.0)
                            doc.add_paragraph("")
                            
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                                cropped.save(tmp.name)
                                pdf.image(tmp.name, w=150)
                            os.unlink(tmp.name)
                            cropped_count += 1
                    except Exception:
                        pass
        
        for p in layout.get("paragraphs", []):
            runs = p.get("runs", [])
            text_join = "".join([r.get("text", "") for r in runs]).strip()
            if not text_join: continue
            
            para = doc.add_paragraph()
            para.alignment = alignment_to_docx(p.get("alignment", "left"))
            
            pdf.set_font("Helvetica", size=12)
            safe_text = text_join.encode('latin-1', 'replace').decode('latin-1')
            
            try:
                pdf.multi_cell(0, 8, text=safe_text)
            except Exception:
                # If a word is too long and causes "Not enough horizontal space", fallback to write
                try:
                    pdf.write(8, text=safe_text + '\n')
                except Exception:
                    pass
            
            for r in runs:
                t = r.get("text", "")
                if not t: continue
                run = para.add_run(t)
                run.bold = bool(r.get("bold", False))
                run.italic = bool(r.get("italic", False))
                
        if idx != len(images): 
            doc.add_page_break()
            
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    
    pdf_io = io.BytesIO(bytes(pdf.output()))
    pdf_io.seek(0)
    return docx_io, pdf_io, cropped_count

# -------------------------
# UI Logic
# -------------------------
st.sidebar.title("Mode Selection")
app_mode = st.sidebar.radio("Choose Version:", ["Phase 1 (Basic Tool)", "Phase 2 (Agentic System)"])

if app_mode == "Phase 1 (Basic Tool)":
    st.title("📄 Image → Formatted DOCX (Phase 1)")
    st.write("Extracts text and formatting from images with a live screen preview.")

    model_id = st.selectbox(
        "Groq Vision Model",
        ["meta-llama/llama-4-scout-17b-16e-instruct", "meta-llama/llama-4-maverick-17b-128e-instruct"],
        key="model_p1"
    )

    include_page_image = st.checkbox("Include original page image", value=True, key="img_p1")

    uploaded = st.file_uploader("Upload images or PDFs", type=["png", "jpg", "jpeg", "webp", "pdf"], accept_multiple_files=True, key="up_p1")

    if uploaded:
        images = process_uploaded_files(uploaded)
        st.divider()

        if st.button("Generate DOCX & PDF", type="primary", key="btn_p1"):
            try:
                with st.spinner("Analyzing layout + generating documents..."):
                    docx_io, pdf_io, _ = process_document(images, model_id, include_page_image)
                
                # --- PREVIEW LOGIC ---
                st.subheader("📝 Live Output Preview")
                docx_io.seek(0)
                result = mammoth.convert_to_html(docx_io)
                html_content = result.value
                
                # Displaying the Word content in a scrollable container
                st.markdown(
                    f'<div style="background-color: white; color: black; padding: 20px; border-radius: 5px; border: 1px solid #ddd; max-height: 400px; overflow-y: auto;">{html_content}</div>', 
                    unsafe_allow_html=True
                )
                
                st.success("Analysis complete.")
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        "⬇ Download DOCX",
                        data=docx_io.getvalue(),
                        file_name="formatted_output_phase1.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="dl_docx_p1"
                    )
                with col2:
                    st.download_button(
                        "⬇ Download PDF",
                        data=pdf_io.getvalue(),
                        file_name="formatted_output_phase1.pdf",
                        mime="application/pdf",
                        key="dl_pdf_p1"
                    )
            except Exception as e:
                st.error(str(e))
    else:
        st.info("Upload images to start.")

else:
    # Phase 2 Logic
    st.title("📄 Intelligent Agent: Image → DOCX (Phase 2)")
    st.write("An autonomous agent that perceives documents, considers ethical constraints (PII), and extracts text/formatting.")

    with st.expander("ℹ️ View Agentic Workflow", expanded=True):
        st.markdown("""
        **1. 👀 Observe:** Analyze uploaded document using Vision Model.
        **2. 🧠 Interpret:** Detect complex diagrams and sensitive PII.
        **3. ⚖️ Decide:** Auto-adjust settings (e.g. preserving original image for diagrams) and require human consent if PII is found.
        **4. ⚡ Act:** Extract text, preserve formatting, and generate final DOCX.
        """)

    model_id = st.selectbox(
        "Groq Vision Model",
        ["meta-llama/llama-4-scout-17b-16e-instruct", "meta-llama/llama-4-maverick-17b-128e-instruct"],
        key="model_p2"
    )

    # User can still override, but Agent will suggest
    include_page_image = st.checkbox("Include original page image (Agent may override)", value=True, key="img_p2")

    uploaded = st.file_uploader("Upload images or PDFs", type=["png", "jpg", "jpeg", "webp", "pdf"], accept_multiple_files=True, key="up_p2")

    if uploaded:
        images = process_uploaded_files(uploaded)
        st.divider()

        if "agent_analysis" not in st.session_state:
            st.session_state.agent_analysis = None

        # Step 1: Agent Perception & Decision
        if st.button("1. Agent: Observe & Decide", type="primary", key="btn_obs_p2"):
            with st.spinner("Agent is observing and interpreting the document..."):
                st.session_state.agent_analysis = groq_agent_perceive(images[0], model_id)
                
        if st.session_state.agent_analysis:
            st.subheader("🤖 Agent Perception & Decision")
            analysis = st.session_state.agent_analysis
            
            st.write(f"**Agent Reasoning:** {analysis.get('reasoning', 'No reasoning provided.')}")
            
            # Decision: Diagrams
            if analysis.get('has_diagrams'):
                st.info("💡 **Decision:** Agent detected diagrams. It will intelligently **crop** and insert only the specific diagrams, keeping the final document clean.")
                include_page_image = False # Disable full page append because we are cropping!
            else:
                st.info("💡 **Decision:** No complex diagrams detected. Text extraction should be sufficient.")
            
            # Ethical/Legal Decision: PII
            if analysis.get('has_pii'):
                st.warning("⚠️ **ETHICAL WARNING (Privacy):** The agent detected potential Personally Identifiable Information (PII) in this document. Processing this data might violate privacy policies.")
                approval = st.checkbox("HUMAN-IN-THE-LOOP: I confirm I have the right to process this PII.", key="chk_pii")
                can_proceed = approval
            else:
                st.success("✅ **Ethical Check:** No PII detected. Data is safe to process.")
                can_proceed = True
                
            if can_proceed:
                # Step 2: Agent Action
                if st.button("2. Agent: Act (Extract & Generate PDF/DOCX)", type="primary", key="btn_act_p2"):
                    try:
                        with st.spinner("Agent is executing text extraction, diagram cropping, and formatting..."):
                            docx_io, pdf_io, cropped_count = process_document(images, model_id, include_page_image)
                        
                        if cropped_count > 0:
                            st.success(f"✂️ Agent successfully cropped and inserted {cropped_count} diagram(s)!")
                        elif analysis.get('has_diagrams') and not include_page_image:
                            st.warning("⚠️ Agent detected a diagram initially, but the OCR extraction failed to map its exact coordinates. You may want to manually check 'Include original page image'.")

                        # --- PREVIEW LOGIC ---
                        st.subheader("📝 Live Output Preview")
                        docx_io.seek(0)
                        result = mammoth.convert_to_html(docx_io)
                        html_content = result.value
                        
                        # Displaying the Word content in a scrollable container
                        st.markdown(
                            f'<div style="background-color: white; color: black; padding: 20px; border-radius: 5px; border: 1px solid #ddd; max-height: 400px; overflow-y: auto;">{html_content}</div>', 
                            unsafe_allow_html=True
                        )
                        
                        st.success("Agent successfully completed the task.")
                        col1, col2 = st.columns(2)
                        with col1:
                            st.download_button(
                                "⬇ Download DOCX",
                                data=docx_io.getvalue(),
                                file_name="formatted_output_agent.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="dl_docx_p2"
                            )
                        with col2:
                            st.download_button(
                                "⬇ Download PDF",
                                data=pdf_io.getvalue(),
                                file_name="formatted_output_agent.pdf",
                                mime="application/pdf",
                                key="dl_pdf_p2"
                            )
                    except Exception as e:
                        st.error(f"Agent execution failed: {str(e)}")
            else:
                st.error("Agent execution paused: Awaiting human override for PII.")
    else:
        st.info("Upload images to start.")